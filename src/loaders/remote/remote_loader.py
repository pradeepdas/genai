# Required packages
# pip install requests PyPDF2 python-docx pyyaml pandas beautifulsoup4

import requests
import PyPDF2
import docx
import json
import yaml
import csv
import pandas as pd
import xml.etree.ElementTree as ET
from io import StringIO, BytesIO
from zipfile import ZipFile
from bs4 import BeautifulSoup
from langchain_core.documents.base import Document
from requests.auth import HTTPBasicAuth

class ContentLoader:
    def __init__(self, url, file_format="auto", auth=None):
        """
        Initialize the ContentLoader.

        Args:
        url (str): The URL to load the content from.
        file_format (str): The format of the file (auto, html, pdf, docx, json, yaml, csv, xml, zip, apple_health).
                           Default is "auto" for automatic detection.
        auth (dict): Dictionary containing authentication method and credentials. Default is None.
                     Example: {'type': 'basic', 'username': 'user', 'password': 'pass'}
        """
        self.url = url
        self.file_format = file_format
        self.auth = self.configure_auth(auth)
        self.content_type = self.detect_content_type() if file_format == "auto" else file_format

    def configure_auth(self, auth):
        if auth:
            if auth['type'] == 'basic':
                return HTTPBasicAuth(auth['username'], auth['password'])
            elif auth['type'] == 'bearer':
                return {'Authorization': f"Bearer {auth['token']}"}
            elif auth['type'] == 'api_key':
                return {'Authorization': f"Token {auth['token']}"}
        return None

    def detect_content_type(self):
        headers = self.auth if isinstance(self.auth, dict) else None
        auth = None if isinstance(self.auth, dict) else self.auth
        response = requests.head(self.url, allow_redirects=True, headers=headers, auth=auth)
        return response.headers.get('content-type')

    def load_content(self):
        try:
            if 'html' in self.content_type:
                content = self.load_web_page()
            elif 'pdf' in self.content_type:
                content = self.load_online_pdf()
            elif 'msword' in self.content_type or 'officedocument.wordprocessingml.document' in self.content_type:
                content = self.load_doc_file()
            elif 'json' in self.content_type or self.content_type == 'json':
                content = self.load_json()
            elif 'yaml' in self.content_type or self.content_type == 'yaml':
                content = self.load_yaml()
            elif 'csv' in self.content_type:
                content = self.load_csv()
            elif 'xml' in self.content_type or self.content_type == 'xml':
                content = self.load_xml()
            elif 'zip' in self.content_type or self.content_type == 'zip':
                content = self.load_zip_file()
            elif 'apple_health' in self.content_type or self.content_type == 'apple_health':
                content = self.load_zip_file(is_apple_health=True)
            else:
                content = "Unsupported file type"
        except Exception as e:
            content = str(e)
        
        document = Document(page_content=content)
        return [document]

    def load_web_page(self):
        headers = self.auth if isinstance(self.auth, dict) else None
        auth = None if isinstance(self.auth, dict) else self.auth
        response = requests.get(self.url, headers=headers, auth=auth)
        soup = BeautifulSoup(response.content, 'html.parser')
        return soup.get_text()

    def load_online_pdf(self):
        headers = self.auth if isinstance(self.auth, dict) else None
        auth = None if isinstance(self.auth, dict) else self.auth
        response = requests.get(self.url, headers=headers, auth=auth)
        with open('temp.pdf', 'wb') as f:
            f.write(response.content)
        with open('temp.pdf', 'rb') as f:
            reader = PyPDF2.PdfFileReader(f)
            text = ''
            for page in range(reader.numPages):
                text += reader.getPage(page).extract_text()
        return text

    def load_doc_file(self):
        headers = self.auth if isinstance(self.auth, dict) else None
        auth = None if isinstance(self.auth, dict) else self.auth
        response = requests.get(self.url, headers=headers, auth=auth)
        with open('temp.docx', 'wb') as f:
            f.write(response.content)
        doc = docx.Document('temp.docx')
        text = ''
        for paragraph in doc.paragraphs:
            text += paragraph.text
        return text

    def load_json(self):
        headers = self.auth if isinstance(self.auth, dict) else None
        auth = None if isinstance(self.auth, dict) else self.auth
        response = requests.get(self.url, headers=headers, auth=auth)
        data = response.json()
        return json.dumps(data)

    def load_yaml(self):
        headers = self.auth if isinstance(self.auth, dict) else None
        auth = None if isinstance(self.auth, dict) else self.auth
        response = requests.get(self.url, headers=headers, auth=auth)
        data = yaml.safe_load(response.text)
        return yaml.dump(data)

    def load_csv(self):
        headers = self.auth if isinstance(self.auth, dict) else None
        auth = None if isinstance(self.auth, dict) else self.auth
        response = requests.get(self.url, headers=headers, auth=auth)
        csvfile = StringIO(response.text)
        df = pd.read_csv(csvfile, parse_dates=True, index_col=0)
        return df.to_csv()

    def load_xml(self):
        headers = self.auth if isinstance(self.auth, dict) else None
        auth = None if isinstance(self.auth, dict) else self.auth
        response = requests.get(self.url, headers=headers, auth=auth)
        tree = ET.ElementTree(ET.fromstring(response.content))
        root = tree.getroot()
        return ET.tostring(root, encoding='unicode')

    def load_zip_file(self, is_apple_health=False):
        headers = self.auth if isinstance(self.auth, dict) else None
        auth = None if isinstance(self.auth, dict) else self.auth
        response = requests.get(self.url, headers=headers, auth=auth)
        zipfile = ZipFile(BytesIO(response.content))
        if is_apple_health:
            return self.load_apple_health_data(zipfile)
        else:
            return self.load_standard_zip(zipfile)

    def load_apple_health_data(self, zipfile):
        health_data = []
        for file in zipfile.namelist():
            if file.endswith('.xml'):
                with zipfile.open(file) as xml_file:
                    tree = ET.parse(xml_file)
                    root = tree.getroot()
                    data = []
                    for record in root.findall('Record'):
                        data.append(record.attrib)
                    health_data.append(data)
        return json.dumps(health_data)

    def load_standard_zip(self, zipfile):
        file_data = {}
        for file in zipfile.namelist():
            with zipfile.open(file) as f:
                file_data[file] = f.read().decode('utf-8')
        return json.dumps(file_data)
